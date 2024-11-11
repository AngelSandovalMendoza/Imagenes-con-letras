import tkinter as tk
from tkinter import filedialog, simpledialog
from PIL import Image, ImageTk
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

CARACTERES_ESPECIALES = [
    ('@', 0, 15), ('#', 16, 31), ('M', 32, 47), ('N', 48, 63),
    ('Q', 64, 79), ('U', 80, 95), ('A', 96, 111), ('D', 112, 127),
    ('0', 128, 143), ('Y', 144, 159), ('2', 160, 175), ('$', 176, 191),
    ('%', 192, 209), ('+', 210, 225), ('.', 226, 239), (' ', 240, 255)
]

MAPA_CARTAS = [
    ('M', 0, 20), ('L', 21, 40), ('K', 41, 60), ('J', 61, 80),
    ('I', 81, 100), ('H', 101, 120), ('G', 121, 140), ('F', 141, 160),
    ('E', 161, 180), ('D', 181, 200), ('C', 201, 220), ('B', 221, 240), ('A', 241, 255)
]


def obtener_caracter_especial(intensidad,diccionario):
    for caracter, min_valor, max_valor in diccionario:
        if min_valor <= intensidad <= max_valor:
            return caracter
    return ' '

def crear_imagen_html(imagen, escala_grises=False):
    ancho, alto = imagen.size
    contenido_html = "<html><body style='background-color: black;'><pre style='font-size: 7px; line-height: 5px;'>\n"

    for y in range(alto):
        for x in range(ancho):
            pixel = imagen.getpixel((x, y))
            if escala_grises:
                valor_gris = int(0.299 * pixel[0] + 0.587 * pixel[1] + 0.114 * pixel[2])
                color = f"rgb({valor_gris},{valor_gris},{valor_gris})"
            else:
                color = f"rgb({pixel[0]},{pixel[1]},{pixel[2]})"
            contenido_html += f"<span style='color: {color};'>M</span>"
        contenido_html += "\n"

    contenido_html += "</pre></body></html>"
    return contenido_html

def crear_imagen_html_con_caracteres_especiales(imagen, escala_grises=False):
    ancho, alto = imagen.size
    contenido_html = "<html><body style='background-color: black;'><pre style='font-size: 7px; line-height: 5px;'>\n"

    for y in range(alto):
        for x in range(ancho):
            pixel = imagen.getpixel((x, y))
            if escala_grises:
                intensidad = int(0.299 * pixel[0] + 0.587 * pixel[1] + 0.114 * pixel[2])
                color = f"rgb({intensidad},{intensidad},{intensidad})"
            else:
                intensidad = int(0.299 * pixel[0] + 0.587 * pixel[1] + 0.114 * pixel[2])
                color = f"rgb({pixel[0]},{pixel[1]},{pixel[2]})"
            caracter = obtener_caracter_especial(intensidad, CARACTERES_ESPECIALES)
            contenido_html += f"<span style='color: {color};'>{caracter}</span>"
        contenido_html += "\n"

    contenido_html += "</pre></body></html>"
    return contenido_html

def crear_docx_con_cartas(imagen, escala_grises=False):
    doc = Document()

    estilo = doc.styles['Normal']
    fuente = estilo.font
    fuente.name = 'Playcrds'
    fuente.size = Pt(3)

    estilo.paragraph_format.space_before = 0
    estilo.paragraph_format.space_after = 0
    estilo.paragraph_format.line_spacing = 1

    ancho, alto = imagen.size
    for y in range(alto):
        parrafo = doc.add_paragraph()

        for x in range(ancho):
            pixel = imagen.getpixel((x, y))
            if escala_grises:
                intensidad = int(0.299 * pixel[0] + 0.587 * pixel[1] + 0.114 * pixel[2])
                color = f"rgb({intensidad},{intensidad},{intensidad})"
            else:
                color = f"rgb({pixel[0]},{pixel[1]},{pixel[2]})"

            caracter = obtener_caracter_especial(intensidad, MAPA_CARTAS)

            run = parrafo.add_run(caracter)
            run.font.size = Pt(3)
            parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph() 

    doc.save('cartas_output.docx')

def crear_imagen_html_con_frase(imagen, frase, escala_grises=False):
    ancho, alto = imagen.size
    contenido_html = "<html><body style='background-color: black;'><pre style='font-size: 7px; line-height: 5px;'>\n"
    frase_len = len(frase)
    
    for y in range(alto):
        for x in range(ancho):
            pixel = imagen.getpixel((x, y))
            if escala_grises:
                intensidad = int(0.299 * pixel[0] + 0.587 * pixel[1] + 0.114 * pixel[2])
                color = f"rgb({intensidad},{intensidad},{intensidad})"
            else:
                color = f"rgb({pixel[0]},{pixel[1]},{pixel[2]})"
            caracter = frase[(y * ancho + x) % frase_len]
            contenido_html += f"<span style='color: {color};'>{caracter}</span>"
        contenido_html += "\n"

    contenido_html += "</pre></body></html>"
    return contenido_html

def guardar_html(imagen, escala_grises=False):
    nombre_archivo = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("Archivos HTML", "*.html")])
    if nombre_archivo:
        contenido_html = crear_imagen_html(imagen, escala_grises)
        with open(nombre_archivo, 'w') as archivo:
            archivo.write(contenido_html)

def guardar_html_con_caracteres_especiales(imagen, escala_grises=False):
    nombre_archivo = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("Archivos HTML", "*.html")])
    if nombre_archivo:
        contenido_html = crear_imagen_html_con_caracteres_especiales(imagen, escala_grises)
        with open(nombre_archivo, 'w') as archivo:
            archivo.write(contenido_html)

def guardar_html_con_frase(imagen, escala_grises=False):
    frase = simpledialog.askstring("Frase personalizada", "Escribe la frase para crear la imagen:")
    if frase:
        nombre_archivo = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("Archivos HTML", "*.html")])
        if nombre_archivo:
            contenido_html = crear_imagen_html_con_frase(imagen, frase, escala_grises)
            with open(nombre_archivo, 'w') as archivo:
                archivo.write(contenido_html)

def abrir_imagen():
    ruta_archivo = filedialog.askopenfilename(filetypes=[("Archivos de imagen", "*.png;*.jpg;*.jpeg;*.bmp")])
    if ruta_archivo:
        imagen = Image.open(ruta_archivo)
        imagen = imagen.resize((100, 100)) 
        etiqueta_imagen.img = ImageTk.PhotoImage(imagen)
        etiqueta_imagen.config(image=etiqueta_imagen.img)
        
        btn_html_color.config(command=lambda: guardar_html(imagen, escala_grises=False), state=tk.NORMAL)
        btn_html_grises.config(command=lambda: guardar_html(imagen, escala_grises=True), state=tk.NORMAL)
        btn_html_color_especial.config(command=lambda: guardar_html_con_caracteres_especiales(imagen, escala_grises=False), state=tk.NORMAL)
        btn_html_grises_especial.config(command=lambda: guardar_html_con_caracteres_especiales(imagen, escala_grises=True), state=tk.NORMAL)
        btn_html_frase_color.config(command=lambda: guardar_html_con_frase(imagen, escala_grises=False), state=tk.NORMAL)
        btn_html_frase_grises.config(command=lambda: guardar_html_con_frase(imagen, escala_grises=True), state=tk.NORMAL)
        btn_cartas.config(command=lambda: crear_docx_con_cartas(imagen, escala_grises=True), state=tk.NORMAL)

# Configuración de la interfaz Tkinter
ventana = tk.Tk()
ventana.title("Generador de HTML con imagen en caracteres especiales o frase personalizada")
ventana.geometry("300x500")

marco = tk.Frame(ventana)
marco.pack(pady=20)

btn_abrir = tk.Button(marco, text="Abrir imagen", command=abrir_imagen)
btn_abrir.pack()

etiqueta_imagen = tk.Label(marco)
etiqueta_imagen.pack()

btn_html_color = tk.Button(marco, text="Guardar HTML (Color con 'M')", state=tk.DISABLED)
btn_html_color.pack(pady=5)

btn_html_grises = tk.Button(marco, text="Guardar HTML (Grises con 'M')", state=tk.DISABLED)
btn_html_grises.pack(pady=5)

btn_html_color_especial = tk.Button(marco, text="Guardar HTML (Color con Especiales)", state=tk.DISABLED)
btn_html_color_especial.pack(pady=5)

btn_html_grises_especial = tk.Button(marco, text="Guardar HTML (Grises con Especiales)", state=tk.DISABLED)
btn_html_grises_especial.pack(pady=5)

btn_html_frase_color = tk.Button(marco, text="Guardar HTML (Color con Frase)", state=tk.DISABLED)
btn_html_frase_color.pack(pady=5)

btn_html_frase_grises = tk.Button(marco, text="Guardar HTML (Grises con Frase)", state=tk.DISABLED)
btn_html_frase_grises.pack(pady=5)

btn_cartas = tk.Button(marco, text="Guardar .doc (Cartas)", state=tk.DISABLED)
btn_cartas.pack(pady=5)

ventana.mainloop()
